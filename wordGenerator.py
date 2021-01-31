import docx
import os
import datetime

doc = docx.Document('mytable.docx')
username = "ThomasTopuz"

c = 0

doc.paragraphs[9].text = username
doc.paragraphs[12].text = datetime.now().strftime("%m/%d/%Y")
doc.paragraphs[16].text = "Settimana dal " +
for i in doc.paragraphs:
    print(i.text, c)
    c += 1

doc.save('mytable.docx')


