import docx
import os
from datetime import datetime
import firebase as fb

doc = docx.Document('../test.docx')
username = "Thomastopuz"

# creazione heading
doc.paragraphs[9].text = username
doc.paragraphs[12].text = datetime.now().strftime("%m/%d/%Y")
doc.paragraphs[16].text = "Settimana dal X al Y"

# ESPE FATTI
# creazione tabella
doc.add_heading('Espe fatti', level=1)
espe_fatti = fb.get_espe_fatti(username)
table_fatti = doc.add_table(rows=1, cols=3)
heading_cells = table_fatti.rows[0].cells
heading_cells[0].text = 'data'
heading_cells[1].text = 'materia'
heading_cells[2].text = 'osservazioni'
table_fatti.style = 'Grid Table 1 Light'

for i in espe_fatti:
    espe_fatto = i.val()
    cells = table_fatti.add_row().cells
    cells[0].text = espe_fatto['data']
    cells[1].text = espe_fatto['materia']
    cells[2].text = ""

doc.add_paragraph()
doc.add_paragraph()
# ESPE RITORNATI
doc.add_heading('Espe Ritornati', level=1)
espe_ritornati = fb.get_espe_ritornati(username)
table_ritornati = doc.add_table(rows=1, cols=5)
table_ritornati.style = 'Grid Table 1 Light'

# heading
heading_cells = table_ritornati.rows[0].cells
heading_cells[0].text = 'data espe'
heading_cells[1].text = 'materia'
heading_cells[2].text = 'nota'
heading_cells[3].text = 'media'
heading_cells[4].text = 'osservazioni'

for i in espe_ritornati:
    espe_ritornato = i.val()
    cells = table_ritornati.add_row().cells
    cells[0].text = espe_ritornato['data']
    cells[1].text = espe_ritornato['materia']
    cells[2].text = espe_ritornato['nota']
    cells[3].text = str(espe_ritornato['media'])
    cells[4].text = ""

doc.save('../test.docx')
