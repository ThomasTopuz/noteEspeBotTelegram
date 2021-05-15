import docx
from datetime import datetime
import src.firebaseconnector.firebase as fb
import os
from src.encryption.encrypter import *


# function to generate a docx to send to the trainer
def genera_docx(username, fullname):
    def get_week_day(n):
        d = datetime.now().strftime("%Y") + "-W" + str(fb.get_week_number())
        r = str(datetime.strptime(d + '-' + n, "%Y-W%W-%w"))
        return r.split(" ", 1)[0]

    filename = "noteEspeSett" + str(datetime.now().isocalendar()[1]) + "_" + fullname + ".docx"
    doc = docx.Document(os.path.join(os.path.dirname(__file__), 'template.docx'))
    # creazione heading
    doc.paragraphs[9].text = fullname
    doc.paragraphs[12].text = datetime.now().strftime("%d.%m.%Y")

    doc.paragraphs[16].text = "Settimana dal " + get_week_day("1") + " al " + get_week_day("5")

    # ESPE FATTI
    # creazione tabella
    doc.add_heading('Espe fatti', level=1)
    espe_fatti = fb.get_espe_fatti(username)
    table_fatti = doc.add_table(rows=1, cols=3)
    heading_cells = table_fatti.rows[0].cells
    heading_cells[0].text = 'data'
    heading_cells[1].text = 'materia'
    heading_cells[2].text = 'osservazioni'
    table_fatti.style = 'Grid Table 1 Light'

    for i in espe_fatti:
        espe_fatto = i.val()
        cells = table_fatti.add_row().cells
        cells[0].text = espe_fatto['data']
        cells[1].text = espe_fatto['materia']
        cells[2].text = decrypt(espe_fatto['osservazioni_fatto'])

    doc.add_paragraph()
    doc.add_paragraph()
    # ESPE RITORNATI
    doc.add_heading('Espe Ritornati', level=1)
    espe_ritornati = fb.get_espe_ritornati(username)
    table_ritornati = doc.add_table(rows=1, cols=5)
    table_ritornati.style = 'Grid Table 1 Light'

    # heading
    heading_cells = table_ritornati.rows[0].cells
    heading_cells[0].text = 'data espe'
    heading_cells[1].text = 'materia'
    heading_cells[2].text = 'nota'
    heading_cells[3].text = 'media'
    heading_cells[4].text = 'osservazioni'

    for i in espe_ritornati:
        espe_ritornato = i.val()
        cells = table_ritornati.add_row().cells
        cells[0].text = espe_ritornato['data']
        cells[1].text = espe_ritornato['materia']
        cells[2].text = decrypt(espe_ritornato['nota'])
        cells[3].text = decrypt(str(espe_ritornato['media']))
        cells[4].text = decrypt(espe_ritornato['osservazioni_ricevuto'])

    # FOOTER
    footer = doc.sections[0].footer
    footer.add_paragraph("Data d'emissione: " + datetime.now().strftime("%m/%d/%Y"))
    doc.save(filename)
    return filename


# function to delete the generated docx once is sent
def elimina_file(filename):
    os.remove(filename)
